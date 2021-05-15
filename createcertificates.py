import os
import openpyxl
from openpyxl.utils import get_column_letter
from openpyxl.utils.cell import get_column_letter
import docx
from datetime import datetime

class CreateCertificates:
    def __init__(self):
        self.intro()
        self.check_files_exist()
        self.try_open_and_assign_excel_spreadsheet()
        self.try_open_and_assign_word_doc()
        self.assign_sheet_class_variable()
        self.check_sheet_has_data()
        self.assign_columns_as_class_variable()
        self.assign_doc_paragraphs_as_class_variable()
        self.create_certificates()
        self.print_completed_message()

    def intro(self):
        print("To proceed we need to find:")
        print("The 'studentcertificates.xlsx' file in the folder 'excelfiles'")
        print("...AND the 'cerfificate.docx' file in the 'wordfiles' folder.")
        print("===Checking folders===")

    def check_files_exist(self):
        self.check_student_certificate_xlsx_file_exists()
        self.check_certificate_docx_file_exists()

    def check_student_certificate_xlsx_file_exists(self):
        os.chdir('excelfiles')
        excel_files = os.listdir()

        if not "studentcertificates.xlsx" in excel_files:
            print("Sorry, we couldn't find the file 'studentcertificates.xlsx' in the excel folder")
            print("Please check and try again")
            quit()

    def check_certificate_docx_file_exists(self):
        os.chdir('../wordfiles')
        word_files = os.listdir()

        if not "certificate.docx" in word_files:
            print("Sorry, we couldn't find the file 'certificate.docx' in the wordfiles folder")
            print("Please check and try again")
            quit()

    def try_open_and_assign_excel_spreadsheet(self):
        os.chdir("../excelfiles")
        try:
            self.excel_spreadsheet = openpyxl.load_workbook('studentcertificates.xlsx')
        except:
            print("We found but couldn't open the file 'studentcertificates.xlsx'")
            quit()

    def try_open_and_assign_word_doc(self):
        os.chdir('../wordfiles')
        try:
            self.certificate_doc = docx.Document('certificate.docx')
        except:
            print("We found but couldn't open the file 'certificate.docx'")
            quit()

    def assign_sheet_class_variable(self):
        self.sheet = self.excel_spreadsheet.active

    def check_sheet_has_data(self):
        if self.sheet.max_row < 2:
            print("Sorry, there aren't any data rows in the 'studentcertificates.xlsx' file. Please check")
            quit()

    def assign_columns_as_class_variable(self):
        columns = list()
        for y in range(1, self.sheet.max_column + 1):
            column_dict = dict()
            column_dict['column'] = get_column_letter(y)
            column_dict['column_name'] = self.sheet[get_column_letter(y) + "1"].value.lower()
            columns.append(column_dict)

        self.columns = columns

    def assign_doc_paragraphs_as_class_variable(self):
        paragraphs = list()
        for para in self.certificate_doc.paragraphs:
            para_dict = dict()
            para_dict['text'] = para.text
            para_dict['font_size'] = para.runs[0].font.size
            para_dict['bold'] = para.runs[0].bold
            paragraphs.append(para_dict)

        self.paragraphs = paragraphs

    def create_certificates(self):
        print("CREATING CERTIFICATES")
        for row_number in range(2, self.sheet.max_row + 1):
            new_certificate = docx.Document()
            self.loop_paragraphs_add_to_new_certificate(row_number, new_certificate)

    def loop_paragraphs_add_to_new_certificate(self, row_number, new_certificate):
        for para in self.paragraphs:
            paragraph_text = self.replace_para_text_with_cell_data(para['text'], row_number)
            paragraph_font = para['font_size']
            paragrpah_bold = para['bold']

            new_para = new_certificate.add_paragraph()
            new_para.text = paragraph_text
            new_para.runs[0].font.size = paragraph_font
            new_para.runs[0].bold = paragrpah_bold
        
        self.save_new_certificate(row_number, new_certificate)

    def replace_para_text_with_cell_data(self, para_text, row_number):
        for column in self.columns:
            if column['column_name'] in para_text:
                para_text = para_text.replace(
                    column['column_name'],
                    str(self.sheet[column['column'] + str(row_number)].value) 
                )
                        
            if 'today' in para_text:
                para_text = para_text.replace(
                    'today',
                    str(datetime.today().strftime('%d/%m/%Y'))
            )

        return para_text

    def save_new_certificate(self, row_number, new_certificate):
        last_name = self.sheet[self.columns[1]['column'] + str(row_number)].value
        todays_date = str(datetime.today().strftime('%d-%m-%Y'))
        new_certificate.save(todays_date + last_name + '.docx')

    def print_completed_message(self):
        print("DONE!")
        print("You can find your certificates for each student created in the 'wordfiles' folder.")
        print("***If your certificates don't look correct, check that the words in the certificate file match the title columns in the speadsheet.***")


